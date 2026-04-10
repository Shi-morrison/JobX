"""Tests for agents/cover_letter.py — Task 3.2
Run with: .venv/bin/pytest tests/test_cover_letter.py -v
"""
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document

from agents.cover_letter import generate_cover_letter, _write_docx
from db.models import Job


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_job(id=1, title="Backend Engineer", company="Stripe", description="We need Python and distributed systems."):
    job = MagicMock(spec=Job)
    job.id = id
    job.title = title
    job.company = company
    job.description = description
    job.gap_analysis = {
        "hard_gaps": ["Kubernetes"],
        "soft_gaps": ["distributed systems"],
        "reframe_suggestions": [
            {"gap": "distributed systems", "suggestion": "Highlight microservices work"}
        ],
    }
    return job


RESUME_DATA = {
    "skills": ["Python", "Go", "React", "PostgreSQL"],
    "experience": [
        {
            "title": "Software Engineer",
            "company": "Corp",
            "start_date": "Jan 2023",
            "end_date": "Present",
            "bullets": ["Built payment APIs", "Led backend team", "Improved latency by 40%"],
        }
    ],
}

FAKE_COVER_LETTER = {
    "subject_line": "Application — Backend Engineer at Stripe",
    "body": "Opening paragraph here.\n\nSecond paragraph about experience.\n\nClosing paragraph.",
    "word_count": 20,
}


# ---------------------------------------------------------------------------
# _write_docx
# ---------------------------------------------------------------------------

def test_write_docx_creates_file(tmp_path):
    output_path = tmp_path / "cover_letter_1.docx"
    _write_docx(output_path, "Application — Engineer at Acme", "Paragraph one.\n\nParagraph two.")
    assert output_path.exists()


def test_write_docx_contains_subject_line(tmp_path):
    output_path = tmp_path / "cover_letter_1.docx"
    _write_docx(output_path, "Application — Backend Engineer at Stripe", "Body text here.")
    doc = Document(str(output_path))
    texts = [p.text for p in doc.paragraphs]
    assert any("Backend Engineer at Stripe" in t for t in texts)


def test_write_docx_splits_on_double_newline(tmp_path):
    output_path = tmp_path / "cl.docx"
    _write_docx(output_path, "Subject", "First paragraph.\n\nSecond paragraph.")
    doc = Document(str(output_path))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "First paragraph." in texts
    assert "Second paragraph." in texts


# ---------------------------------------------------------------------------
# generate_cover_letter
# ---------------------------------------------------------------------------

def test_generate_cover_letter_returns_file_path(tmp_path):
    with patch("agents.cover_letter.OUTPUT_DIR", tmp_path), \
         patch("agents.cover_letter.ClaudeClient") as MockClient, \
         patch("agents.cover_letter.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = FAKE_COVER_LETTER
        result = generate_cover_letter(_make_job(), RESUME_DATA)

    assert "cover_letter_1.docx" in result["file_path"]
    assert Path(result["file_path"]).exists()


def test_generate_cover_letter_returns_subject_and_wordcount(tmp_path):
    with patch("agents.cover_letter.OUTPUT_DIR", tmp_path), \
         patch("agents.cover_letter.ClaudeClient") as MockClient, \
         patch("agents.cover_letter.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = FAKE_COVER_LETTER
        result = generate_cover_letter(_make_job(), RESUME_DATA)

    assert result["subject_line"] == "Application — Backend Engineer at Stripe"
    assert result["word_count"] == 20


def test_generate_cover_letter_fallback_word_count(tmp_path):
    """If word_count is missing from Claude response, fall back to counting words."""
    fake = {**FAKE_COVER_LETTER, "word_count": None}
    with patch("agents.cover_letter.OUTPUT_DIR", tmp_path), \
         patch("agents.cover_letter.ClaudeClient") as MockClient, \
         patch("agents.cover_letter.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = fake
        result = generate_cover_letter(_make_job(), RESUME_DATA)

    assert result["word_count"] > 0


def test_generate_cover_letter_uses_job_fields_in_prompt(tmp_path):
    """JD, company name, and gap info should all be passed into the prompt."""
    captured = {}

    def fake_load_prompt(template_name, **kwargs):
        captured.update(kwargs)
        return "prompt"

    with patch("agents.cover_letter.OUTPUT_DIR", tmp_path), \
         patch("agents.cover_letter.ClaudeClient") as MockClient, \
         patch("agents.cover_letter.load_prompt", side_effect=fake_load_prompt):
        MockClient.return_value.chat_json.return_value = FAKE_COVER_LETTER
        generate_cover_letter(_make_job(), RESUME_DATA)

    assert captured["company"] == "Stripe"
    assert captured["job_title"] == "Backend Engineer"
    assert "Kubernetes" in captured["hard_gaps"]
    assert "microservices" in captured["reframe_suggestions"]


def test_generate_cover_letter_fallback_subject(tmp_path):
    """If Claude omits subject_line, fall back to default."""
    fake = {**FAKE_COVER_LETTER, "subject_line": ""}
    with patch("agents.cover_letter.OUTPUT_DIR", tmp_path), \
         patch("agents.cover_letter.ClaudeClient") as MockClient, \
         patch("agents.cover_letter.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = fake
        result = generate_cover_letter(_make_job(), RESUME_DATA)

    assert "Backend Engineer" in result["subject_line"]
    assert "Stripe" in result["subject_line"]
