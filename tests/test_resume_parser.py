"""Tests for the resume parser in tools/llm.py — run with: .venv/bin/pytest tests/test_resume_parser.py -v"""
import json
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock


# ---------------------------------------------------------------------------
# _extract_resume_text
# ---------------------------------------------------------------------------

def test_extract_resume_text_missing_file(monkeypatch, tmp_path):
    monkeypatch.setattr("tools.llm._RESUME_DOCX", tmp_path / "missing.docx")
    from tools.llm import _extract_resume_text
    with pytest.raises(FileNotFoundError, match="Resume not found"):
        _extract_resume_text()


def test_extract_resume_text_reads_paragraphs(tmp_path, monkeypatch):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Shian Morrison")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("")  # blank — should be skipped
    docx_path = tmp_path / "resume.docx"
    doc.save(str(docx_path))

    monkeypatch.setattr("tools.llm._RESUME_DOCX", docx_path)
    from tools.llm import _extract_resume_text
    text = _extract_resume_text()
    assert "Shian Morrison" in text
    assert "Software Engineer" in text
    assert text.count("\n\n") == 0  # blank lines stripped


# ---------------------------------------------------------------------------
# parse_resume
# ---------------------------------------------------------------------------

def test_parse_resume_uses_cache(tmp_path, monkeypatch):
    """If cache file exists, Claude should NOT be called."""
    cache = tmp_path / "resume_parsed.json"
    cached_data = {"name": "Cached Name", "skills": ["Python"]}
    cache.write_text(json.dumps(cached_data))

    monkeypatch.setattr("tools.llm._RESUME_CACHE_FILE", cache)

    with patch("tools.llm.ClaudeClient") as mock_claude:
        from tools.llm import parse_resume
        result = parse_resume(force=False)
        mock_claude.assert_not_called()

    assert result["name"] == "Cached Name"


def test_parse_resume_force_bypasses_cache(tmp_path, monkeypatch):
    """force=True should call Claude even if cache exists."""
    cache = tmp_path / "resume_parsed.json"
    cache.write_text(json.dumps({"name": "Old Cache"}))

    monkeypatch.setattr("tools.llm._RESUME_CACHE_FILE", cache)

    fake_result = {"name": "Fresh Parse", "skills": ["Go", "Python"], "experience": [],
                   "education": [], "projects": [], "contact": {}}

    with patch("tools.llm._extract_resume_text", return_value="resume text"), \
         patch("tools.llm.load_prompt", return_value="prompt"), \
         patch("tools.llm.ClaudeClient") as MockClient:
        MockClient.return_value.chat_json.return_value = fake_result
        from tools.llm import parse_resume
        result = parse_resume(force=True)

    assert result["name"] == "Fresh Parse"


def test_parse_resume_writes_cache(tmp_path, monkeypatch):
    """After a successful parse, result should be written to cache."""
    cache = tmp_path / "resume_parsed.json"
    monkeypatch.setattr("tools.llm._RESUME_CACHE_FILE", cache)

    fake_result = {"name": "Test User", "skills": ["Python"], "experience": [],
                   "education": [], "projects": [], "contact": {}}

    with patch("tools.llm._extract_resume_text", return_value="resume text"), \
         patch("tools.llm.load_prompt", return_value="prompt"), \
         patch("tools.llm.ClaudeClient") as MockClient:
        MockClient.return_value.chat_json.return_value = fake_result
        from tools.llm import parse_resume
        parse_resume(force=True)

    assert cache.exists()
    saved = json.loads(cache.read_text())
    assert saved["name"] == "Test User"
