"""Tests for agents/resume_tailor.py — Task 3.1
Run with: .venv/bin/pytest tests/test_resume_tailor.py -v
"""
import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock
from docx import Document

from agents.resume_tailor import tailor_resume, _extract_bullets
from db.models import Job


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_job(id=1, title="Backend Engineer", company="Acme", description="We need Python and Kubernetes."):
    job = MagicMock(spec=Job)
    job.id = id
    job.title = title
    job.company = company
    job.description = description
    return job


def _make_docx_with_bullets(tmp_path: Path, bullets: list[str]) -> Path:
    """Create a temporary .docx with given bullet paragraphs."""
    doc = Document()
    doc.add_paragraph("John Doe")           # short — should be skipped
    doc.add_paragraph("EXPERIENCE")         # all-caps — should be skipped
    for bullet in bullets:
        doc.add_paragraph(bullet)
    path = tmp_path / "base_resume.docx"
    doc.save(str(path))
    return path


GAP_ANALYSIS = {
    "hard_gaps": ["Kubernetes"],
    "soft_gaps": ["distributed systems"],
    "reframe_suggestions": [
        {"gap": "distributed systems", "suggestion": "Highlight your microservices work at Corp"}
    ],
}

RESUME_DATA = {
    "skills": ["Python", "Go", "React"],
    "experience": [
        {
            "title": "SWE",
            "company": "Corp",
            "start_date": "2023",
            "end_date": "Present",
            "bullets": ["Built REST APIs in Python serving 10M requests/day"],
        }
    ],
}


# ---------------------------------------------------------------------------
# _extract_bullets
# ---------------------------------------------------------------------------

def test_extract_bullets_skips_short_lines(tmp_path):
    path = _make_docx_with_bullets(tmp_path, ["Built microservices in Python and Go for high-traffic data pipelines"])
    doc = Document(str(path))
    bullets = _extract_bullets(doc)
    assert "John Doe" not in bullets
    assert any("microservices" in b for b in bullets)


def test_extract_bullets_skips_all_caps(tmp_path):
    path = _make_docx_with_bullets(tmp_path, ["Led backend platform team across three product lines"])
    doc = Document(str(path))
    bullets = _extract_bullets(doc)
    assert "EXPERIENCE" not in bullets


def test_extract_bullets_returns_long_paragraphs(tmp_path):
    long_bullet = "Designed and implemented a distributed caching layer that reduced database load by 40%"
    path = _make_docx_with_bullets(tmp_path, [long_bullet])
    doc = Document(str(path))
    bullets = _extract_bullets(doc)
    assert long_bullet in bullets


# ---------------------------------------------------------------------------
# tailor_resume
# ---------------------------------------------------------------------------

def test_tailor_resume_applies_rewrites(tmp_path):
    original = "Built REST APIs in Python serving 10M requests/day"
    rewritten = "Built distributed REST APIs in Python and Go serving 10M requests/day"

    fake_result = {
        "rewrites": [{"original": original, "rewritten": rewritten}],
        "changes_summary": "Emphasized distributed systems experience.",
    }

    resume_path = _make_docx_with_bullets(tmp_path, [original])

    with patch("agents.resume_tailor.BASE_RESUME_PATH", resume_path), \
         patch("agents.resume_tailor.OUTPUT_DIR", tmp_path), \
         patch("agents.resume_tailor.ClaudeClient") as MockClient, \
         patch("agents.resume_tailor.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = fake_result
        result = tailor_resume(_make_job(), RESUME_DATA, GAP_ANALYSIS)

    assert result["rewrites_applied"] == 1
    assert result["total_suggested"] == 1
    assert result["changes_summary"] == "Emphasized distributed systems experience."

    # Verify the saved docx contains the rewritten text
    saved_doc = Document(result["file_path"])
    texts = [p.text for p in saved_doc.paragraphs]
    assert rewritten in texts
    assert original not in texts


def test_tailor_resume_saves_to_correct_path(tmp_path):
    fake_result = {
        "rewrites": [],
        "changes_summary": "No changes needed.",
    }
    resume_path = _make_docx_with_bullets(tmp_path, ["Built backend systems in Go for financial data pipelines"])

    with patch("agents.resume_tailor.BASE_RESUME_PATH", resume_path), \
         patch("agents.resume_tailor.OUTPUT_DIR", tmp_path), \
         patch("agents.resume_tailor.ClaudeClient") as MockClient, \
         patch("agents.resume_tailor.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = fake_result
        result = tailor_resume(_make_job(id=42), RESUME_DATA, GAP_ANALYSIS)

    assert "resume_42.docx" in result["file_path"]


def test_tailor_resume_raises_if_no_base_resume(tmp_path):
    missing_path = tmp_path / "nonexistent.docx"
    with patch("agents.resume_tailor.BASE_RESUME_PATH", missing_path):
        with pytest.raises(FileNotFoundError, match="Base resume not found"):
            tailor_resume(_make_job(), RESUME_DATA, GAP_ANALYSIS)


def test_tailor_resume_no_rewrites_still_saves(tmp_path):
    """If Claude suggests no rewrites, the file should still be saved."""
    fake_result = {"rewrites": [], "changes_summary": "Resume already well-aligned."}
    resume_path = _make_docx_with_bullets(tmp_path, ["Designed APIs in Python used by enterprise clients globally"])

    with patch("agents.resume_tailor.BASE_RESUME_PATH", resume_path), \
         patch("agents.resume_tailor.OUTPUT_DIR", tmp_path), \
         patch("agents.resume_tailor.ClaudeClient") as MockClient, \
         patch("agents.resume_tailor.load_prompt", return_value="prompt"):
        MockClient.return_value.chat_json.return_value = fake_result
        result = tailor_resume(_make_job(), RESUME_DATA, GAP_ANALYSIS)

    assert result["rewrites_applied"] == 0
    assert Path(result["file_path"]).exists()


def test_tailor_resume_uses_reframe_suggestions_in_prompt(tmp_path):
    """Reframe suggestions should be passed into the prompt."""
    captured = {}

    def fake_load_prompt(template_name, **kwargs):
        captured.update(kwargs)
        return "prompt"

    fake_result = {"rewrites": [], "changes_summary": ""}
    resume_path = _make_docx_with_bullets(tmp_path, ["Built and shipped Python microservices for payments platform"])

    with patch("agents.resume_tailor.BASE_RESUME_PATH", resume_path), \
         patch("agents.resume_tailor.OUTPUT_DIR", tmp_path), \
         patch("agents.resume_tailor.ClaudeClient") as MockClient, \
         patch("agents.resume_tailor.load_prompt", side_effect=fake_load_prompt):
        MockClient.return_value.chat_json.return_value = fake_result
        tailor_resume(_make_job(), RESUME_DATA, GAP_ANALYSIS)

    assert "distributed systems" in captured["reframe_suggestions"]
    assert "microservices" in captured["reframe_suggestions"]
