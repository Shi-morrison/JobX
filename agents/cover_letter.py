from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console

from db.session import get_session
from db.models import Job
from tools.llm import ClaudeClient, load_prompt, parse_resume
from agents.scorer import _build_experience_summary

console = Console()

OUTPUT_DIR = Path("data/cover_letters")


def generate_cover_letter(
    job: Job,
    resume_data: dict,
    company_research: dict | None = None,
) -> dict:
    """Generate a tailored cover letter for the job and save as .docx.

    Args:
        job: The job to write the cover letter for (must be scored).
        resume_data: Parsed resume dict from parse_resume().
        company_research: Optional company intel from Phase 4 — not yet used,
                          reserved for when company research agent is built.

    Returns:
        Dict with file_path, subject_line, word_count.
    """
    skills = ", ".join(resume_data.get("skills", []))
    experience_summary = _build_experience_summary(resume_data)
    gap_analysis = job.gap_analysis or {}
    description = (job.description or "")[:3000]

    hard_gaps = ", ".join(gap_analysis.get("hard_gaps", [])) or "None"
    soft_gaps = ", ".join(gap_analysis.get("soft_gaps", [])) or "None"
    reframe_suggestions = gap_analysis.get("reframe_suggestions", [])
    reframe_text = "\n".join(
        f"- {s['gap']}: {s['suggestion']}" for s in reframe_suggestions
    ) or "None"

    # Build company intel context from research if available
    company_intel = "No company research available."
    if company_research:
        parts = []
        if company_research.get("summary"):
            parts.append(company_research["summary"])
        if company_research.get("funding_stage"):
            parts.append(f"Stage: {company_research['funding_stage']}")
        if company_research.get("tech_stack"):
            parts.append(f"Tech stack: {', '.join(company_research['tech_stack'][:10])}")
        if parts:
            company_intel = " | ".join(parts)

    prompt = load_prompt(
        "cover_letter",
        job_title=job.title,
        company=job.company,
        job_description=description,
        skills=skills,
        experience_summary=experience_summary,
        hard_gaps=hard_gaps,
        soft_gaps=soft_gaps,
        reframe_suggestions=reframe_text,
        company_intel=company_intel,
    )

    client = ClaudeClient()
    result = client.chat_json(
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1024,
    )

    body = result.get("body", "")
    subject_line = result.get("subject_line") or f"Application — {job.title} at {job.company}"
    word_count = result.get("word_count") or len(body.split())

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"cover_letter_{job.id}.docx"
    _write_docx(output_path, subject_line, body)

    return {
        "file_path": str(output_path),
        "subject_line": subject_line,
        "word_count": word_count,
    }


def _write_docx(output_path: Path, subject_line: str, body: str) -> None:
    """Write the cover letter text to a formatted .docx file."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Subject / header line
    subj_para = doc.add_paragraph()
    subj_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subj_para.add_run(subject_line)
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()  # blank spacer

    # Body — split on double newlines (paragraph breaks)
    for paragraph_text in body.split("\n\n"):
        text = paragraph_text.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.size = Pt(11)

    doc.save(str(output_path))


def run_cover_letter(job_id: int) -> None:
    """Generate a cover letter for a specific job.

    Requires the job to have been scored first (uses gap_analysis for tone/reframes).
    """
    resume_data = parse_resume()

    with get_session() as db:
        job = db.query(Job).filter(Job.id == job_id).first()

    if not job:
        console.print(f"[red]No job found with ID {job_id}. Find IDs with: python main.py jobs[/red]")
        return

    if not job.fit_score:
        console.print(
            f"[yellow]Job {job_id} has not been scored yet. "
            f"Run: python main.py score --job-id {job_id}[/yellow]"
        )
        return

    # Auto-load cached company research if available
    from agents.company_research import get_cached_research
    company_research = get_cached_research(job.company)
    if company_research:
        console.print(f"[dim]Using cached company research for {job.company}.[/dim]")

    console.print(f"[dim]Generating cover letter for {job.title} @ {job.company}...[/dim]")

    with console.status("Writing cover letter..."):
        try:
            result = generate_cover_letter(job, resume_data, company_research=company_research)
        except Exception as e:
            console.print(f"[red]Cover letter generation failed: {e}[/red]")
            return

    console.print(f"[green]Cover letter saved:[/green] {result['file_path']}")
    console.print(f"[dim]Subject: {result['subject_line']}[/dim]")
    console.print(f"[dim]Word count: {result['word_count']}[/dim]")
    console.print(f"[dim]Open in Word or VS Code to review.[/dim]")
