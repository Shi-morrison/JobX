from pathlib import Path

from docx import Document
from rich.console import Console

from db.session import get_session
from db.models import Job, ResumeVersion
from tools.llm import ClaudeClient, load_prompt, parse_resume

console = Console()

BASE_RESUME_PATH = Path("data/base_resume.docx")
OUTPUT_DIR = Path("data/resume_versions")


def _extract_bullets(doc: Document) -> list[str]:
    """Extract paragraphs that are likely bullet points (experience descriptions).

    Skips very short lines (headers, dates, names) and all-caps section headers.
    """
    bullets = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if len(text) < 30:
            continue
        if text.isupper():
            continue  # section headers like "EXPERIENCE", "SKILLS"
        style_name = (para.style.name or "").lower()
        if any(s in style_name for s in ["list", "bullet", "normal"]):
            bullets.append(text)
    return bullets


def tailor_resume(job: Job, resume_data: dict, gap_analysis: dict) -> dict:
    """Rewrite resume bullets to match the job description.

    Reads data/base_resume.docx, asks Claude to rewrite specific bullets
    to match the JD using reframe_suggestions from gap analysis, and saves
    the result to data/resume_versions/resume_{job_id}.docx.

    Returns:
        Dict with file_path, changes_summary, rewrites_applied, total_suggested.
    """
    if not BASE_RESUME_PATH.exists():
        raise FileNotFoundError(
            f"Base resume not found at {BASE_RESUME_PATH}. "
            "Add your resume at data/base_resume.docx first."
        )

    doc = Document(str(BASE_RESUME_PATH))
    bullets = _extract_bullets(doc)

    if not bullets:
        raise ValueError(
            "No bullet points found in base resume. "
            "Make sure your resume uses standard paragraph styles."
        )

    reframe_suggestions = gap_analysis.get("reframe_suggestions", [])
    reframe_text = "\n".join(
        f"- {s['gap']}: {s['suggestion']}" for s in reframe_suggestions
    ) or "No specific reframes needed — strengthen general alignment with the JD."

    description = (job.description or "")[:3000]

    prompt = load_prompt(
        "tailor_resume",
        job_title=job.title,
        company=job.company,
        job_description=description,
        current_bullets="\n".join(f"- {b}" for b in bullets),
        reframe_suggestions=reframe_text,
    )

    client = ClaudeClient()
    result = client.chat_json(
        messages=[{"role": "user", "content": prompt}],
        max_tokens=2048,
    )

    rewrites = result.get("rewrites", [])
    changes_summary = result.get("changes_summary", "")

    # Build lookup: original text → rewritten text
    rewrite_map = {r["original"].strip(): r["rewritten"] for r in rewrites}

    # Apply rewrites to a fresh copy of the base resume
    output_doc = Document(str(BASE_RESUME_PATH))
    applied = 0

    for para in output_doc.paragraphs:
        text = para.text.strip()
        if text in rewrite_map:
            new_text = rewrite_map[text]
            if para.runs:
                # Preserve run formatting: put all new text in first run, clear rest
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            applied += 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"resume_{job.id}.docx"
    output_doc.save(str(output_path))

    return {
        "file_path": str(output_path),
        "changes_summary": changes_summary,
        "rewrites_applied": applied,
        "total_suggested": len(rewrites),
    }


def run_tailor(job_id: int) -> None:
    """Tailor the base resume for a specific job and save the result.

    Requires the job to have been scored first (needs gap_analysis).
    """
    resume_data = parse_resume()

    with get_session() as db:
        job = db.query(Job).filter(Job.id == job_id).first()

    if not job:
        console.print(f"[red]No job found with ID {job_id}. Find IDs with: python main.py jobs[/red]")
        return

    if not job.fit_score:
        console.print(
            f"[yellow]Job {job_id} has not been scored yet. "
            f"Run: python main.py score --job-id {job_id}[/yellow]"
        )
        return

    gap_analysis = job.gap_analysis or {}

    console.print(f"[dim]Tailoring resume for {job.title} @ {job.company}...[/dim]")

    with console.status("Rewriting bullets to match JD..."):
        try:
            result = tailor_resume(job, resume_data, gap_analysis)
        except Exception as e:
            console.print(f"[red]Resume tailoring failed: {e}[/red]")
            return

    # Save ResumeVersion record to DB
    with get_session() as db:
        rv = ResumeVersion(
            job_id=job_id,
            file_path=result["file_path"],
            changes_summary=result["changes_summary"],
        )
        db.add(rv)

    console.print(f"[green]Resume saved:[/green] {result['file_path']}")
    console.print(
        f"[dim]{result['rewrites_applied']} of {result['total_suggested']} "
        f"bullets rewritten[/dim]"
    )
    if result["changes_summary"]:
        console.print(f"\n[bold]Changes summary:[/bold]")
        console.print(f"  {result['changes_summary']}")
    console.print(
        f"\n[dim]Open in Word or VS Code to review. "
        f"Base resume unchanged at {BASE_RESUME_PATH}[/dim]"
    )
